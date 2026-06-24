"""
DOCX 解析器
============
解析 Microsoft Word (.docx) 格式文档。
"""

import logging
from pathlib import Path

from services.parsers.base import BaseParser, ParseError, ParseResult

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    """Word 文档解析器 — 只负责 .docx 格式。

    提取内容：
    - 段落文本
    - 表格内容（转为管道符分隔文本）

    遵循 Single Responsibility 原则。
    """

    @property
    def format(self) -> str:
        """返回格式标识符。"""
        return "docx"

    def do_parse(self, file_path: Path) -> ParseResult:
        """解析 .docx 文件。

        Args:
            file_path: 文件路径

        Returns:
            ParseResult — content 为提取的完整文本

        Raises:
            ParseError: 文件损坏或无法解析
        """
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise ParseError(
                message="缺少 python-docx 依赖，请执行: pip install python-docx",
                file_path=file_path,
                cause=exc,
            )

        try:
            doc = DocxDocument(str(file_path))
        except Exception as exc:
            raise ParseError(
                message=f"DOCX 文件打开失败，文件可能已损坏: {exc}",
                file_path=file_path,
                cause=exc,
            )

        parts: list[str] = []

        # ── 提取段落文本 ────────────────────────
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # ── 提取表格内容 ────────────────────────
        for table in doc.tables:
            for row in table.rows:
                cells = [
                    cell.text.strip().replace("\n", " ")
                    for cell in row.cells
                ]
                row_text = " | ".join(c for c in cells if c)
                if row_text:
                    parts.append(row_text)

        content = "\n\n".join(parts)

        if not content.strip():
            raise ParseError(
                message="DOCX 文件中未提取到文本内容",
                file_path=file_path,
            )

        logger.info(
            "DOCX 解析完成: %s → %d 段落/行, %d 字符",
            file_path.name,
            len(parts),
            len(content),
        )

        return ParseResult(
            content=content,
            filename=file_path.name,
            format="docx",
        )
