"""
文档解析器 — 多格式文档解析与分块
====================================
支持 .txt / .md / .yaml / .yml / .docx / .pdf 格式。
对长文档实现自动分块（Chunk），便于后续 AI 分段处理。
"""

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Protocol

from config import CHUNK_SIZE, CHUNK_OVERLAP, SUPPORTED_EXTENSIONS

logger = logging.getLogger(__name__)


# ══════════════════════════════════════════════════════════
# 数据模型
# ══════════════════════════════════════════════════════════

@dataclass
class DocumentChunk:
    """文档块。

    Attributes:
        index: 块的序号（从 0 开始）
        content: 块的文本内容
        start_char: 在原文中的起始字符位置
        end_char: 在原文中的结束字符位置
    """
    index: int
    content: str
    start_char: int
    end_char: int


@dataclass
class ParsedDocument:
    """解析后的文档。

    Attributes:
        filename: 原始文件名
        file_type: 文件扩展名
        full_text: 完整文本内容
        char_count: 总字符数
        chunks: 文档分块列表
        page_count: PDF 页数（仅 PDF 有值）
    """
    filename: str
    file_type: str
    full_text: str
    char_count: int
    chunks: list[DocumentChunk] = field(default_factory=list)
    page_count: int | None = None


# ══════════════════════════════════════════════════════════
# 解析器协议
# ══════════════════════════════════════════════════════════

class BaseParser(Protocol):
    """文档解析器协议 — 未来新增格式只需实现此协议。"""

    def parse(self, file_path: Path) -> str:
        """解析文件，返回纯文本内容。"""
        ...


# ══════════════════════════════════════════════════════════
# 各格式解析实现
# ══════════════════════════════════════════════════════════

def _parse_txt(file_path: Path) -> str:
    """解析纯文本文件。"""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _parse_md(file_path: Path) -> str:
    """解析 Markdown 文件（本质是纯文本）。"""
    return _parse_txt(file_path)


def _parse_docx(file_path: Path) -> str:
    """解析 Word (.docx) 文件。

    提取段落文本和表格内容。
    """
    from docx import Document as DocxDocument

    doc = DocxDocument(str(file_path))
    paragraphs: list[str] = []

    # ── 提取段落文本 ────────────────────────────
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # ── 提取表格内容 ────────────────────────────
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join(c for c in cells if c)
            if row_text:
                paragraphs.append(row_text)

    return "\n\n".join(paragraphs)


def _parse_pdf(file_path: Path) -> tuple[str, int]:
    """解析 PDF 文件。

    返回 (纯文本内容, 页数)。
    """
    from PyPDF2 import PdfReader

    reader = PdfReader(str(file_path))
    pages: list[str] = []

    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text.strip())

    return "\n\n".join(pages), len(reader.pages)


# ══════════════════════════════════════════════════════════
# 文档解析器类
# ══════════════════════════════════════════════════════════

class DocumentParser:
    """文档解析器 — 自动识别格式并解析，支持长文档分块。

    使用方式:
        parser = DocumentParser()
        doc = parser.parse("需求文档.docx")
        for chunk in doc.chunks:
            process(chunk.content)
    """

    # 支持的格式与对应的解析函数
    _PARSERS = {
        ".txt": _parse_txt,
        ".md": _parse_md,
        ".yaml": _parse_txt,
        ".yml": _parse_txt,
        ".docx": _parse_docx,
    }

    def __init__(
        self,
        chunk_size: int = CHUNK_SIZE,
        chunk_overlap: int = CHUNK_OVERLAP,
    ) -> None:
        """初始化解析器。

        Args:
            chunk_size: 每个 chunk 的最大字符数
            chunk_overlap: 相邻 chunk 的重叠字符数
        """
        self._chunk_size = chunk_size
        self._chunk_overlap = chunk_overlap

    def parse(self, file_path: str | Path) -> ParsedDocument:
        """解析文档并自动分块。

        Args:
            file_path: 文件路径

        Returns:
            ParsedDocument: 包含完整文本和分块信息的解析结果

        Raises:
            ValueError: 不支持的文件格式
            FileNotFoundError: 文件不存在
        """
        file_path = Path(file_path)

        ext = file_path.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"不支持的文件格式: {ext}，"
                f"仅支持: {', '.join(SUPPORTED_EXTENSIONS)}"
            )

        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        logger.info("开始解析文档: %s (格式: %s)", file_path.name, ext)

        # ── 根据格式选择解析方式 ──────────────────
        page_count: int | None = None

        if ext == ".pdf":
            full_text, page_count = _parse_pdf(file_path)
        elif ext in self._PARSERS:
            full_text = self._PARSERS[ext](file_path)
        else:
            raise ValueError(f"未注册的格式: {ext}")

        # ── 文本清洗 ────────────────────────────────
        full_text = self._clean_text(full_text)

        # ── 分块处理 ────────────────────────────────
        chunks = self._chunk_text(full_text)

        logger.info(
            "文档解析完成: %s → 总字符=%d, 分块数=%d, 页数=%s",
            file_path.name,
            len(full_text),
            len(chunks),
            page_count or "N/A",
        )

        return ParsedDocument(
            filename=file_path.name,
            file_type=ext,
            full_text=full_text,
            char_count=len(full_text),
            chunks=chunks,
            page_count=page_count,
        )

    def _clean_text(self, text: str) -> str:
        """清洗文本：去除多余空行、统一换行符。

        Args:
            text: 原始文本

        Returns:
            清洗后的文本
        """
        # 统一换行符
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        # 去除连续 3 个以上的换行，保留 2 个
        import re
        text = re.sub(r"\n{3,}", "\n\n", text)
        # 去除行内多余空格
        text = re.sub(r"[ \t]{2,}", " ", text)
        return text.strip()

    def _chunk_text(self, text: str) -> list[DocumentChunk]:
        """将长文本切分为重叠的块。

        分块策略：
        - 优先在段落边界切分（遇到两个换行符）
        - 如果段落仍超过 chunk_size，在句子边界切分（遇到句号+空格）
        - 实在不行才硬切分

        Args:
            text: 要分块的文本

        Returns:
            DocumentChunk 列表
        """
        if len(text) <= self._chunk_size:
            return [DocumentChunk(
                index=0,
                content=text,
                start_char=0,
                end_char=len(text),
            )]

        chunks: list[DocumentChunk] = []
        index = 0
        start = 0

        while start < len(text):
            end = start + self._chunk_size

            if end >= len(text):
                # 最后一个块
                chunks.append(DocumentChunk(
                    index=index,
                    content=text[start:].strip(),
                    start_char=start,
                    end_char=len(text),
                ))
                break

            # ── 在段落边界切分 ────────────────────
            # 在 chunk_size 范围内找最后一个双换行符
            chunk_text = text[start:end]
            last_para = chunk_text.rfind("\n\n")
            if last_para > self._chunk_size // 2:
                end = start + last_para
                chunks.append(DocumentChunk(
                    index=index,
                    content=text[start:end].strip(),
                    start_char=start,
                    end_char=end,
                ))
                index += 1
                start = max(start, end - self._chunk_overlap)
                continue

            # ── 在句子边界切分 ────────────────────
            last_period = chunk_text.rfind("。")
            if last_period > self._chunk_size // 2:
                end = start + last_period + 1
                chunks.append(DocumentChunk(
                    index=index,
                    content=text[start:end].strip(),
                    start_char=start,
                    end_char=end,
                ))
                start = max(end - self._chunk_overlap, start + 1)
                index += 1
                continue

            # ── 硬切分 ────────────────────────────
            chunks.append(DocumentChunk(
                index=index,
                content=text[start:end].strip(),
                start_char=start,
                end_char=end,
            ))
            start = end - self._chunk_overlap
            index += 1

        return chunks


# ══════════════════════════════════════════════════════════
# 便捷函数
# ══════════════════════════════════════════════════════════

def parse_document(file_path: str | Path) -> ParsedDocument:
    """解析文档的便捷函数（使用默认配置）。

    Args:
        file_path: 文件路径

    Returns:
        ParsedDocument 对象
    """
    parser = DocumentParser()
    return parser.parse(file_path)
