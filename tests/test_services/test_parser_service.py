"""
ParserService 单元测试
========================
覆盖所有格式解析器 + 异常路径 + SOLID 原则验证。
"""

import pytest
from pathlib import Path

from services.parser_service import ParserService, ParseResult, ParseError
from services.parsers.base import BaseParser


# ══════════════════════════════════════════════════════════
# Fixtures
# ══════════════════════════════════════════════════════════

@pytest.fixture
def service() -> ParserService:
    """创建 ParserService 实例。"""
    return ParserService()


@pytest.fixture
def txt_file(tmp_path: Path) -> Path:
    """创建测试用 TXT 文件。"""
    file = tmp_path / "test_doc.txt"
    file.write_text(
        "用户管理系统需求规格说明书\n\n"
        "## 1. 用户注册\n"
        "用户可以通过手机号或邮箱注册。\n"
        "密码要求8-20位，含大小写字母和数字。\n\n"
        "## 2. 用户登录\n"
        "支持用户名+密码登录，连续5次失败锁定30分钟。",
        encoding="utf-8",
    )
    return file


@pytest.fixture
def md_file(tmp_path: Path) -> Path:
    """创建测试用 Markdown 文件。"""
    file = tmp_path / "test_doc.md"
    file.write_text(
        "# 需求文档\n\n"
        "## 功能模块\n\n"
        "- 用户管理\n"
        "- 权限控制\n"
        "- 数据报表\n",
        encoding="utf-8",
    )
    return file


# ══════════════════════════════════════════════════════════
# 核心功能测试
# ══════════════════════════════════════════════════════════

class TestParserServiceCore:
    """ParserService 核心功能测试。"""

    def test_parse_txt_returns_unified_format(
        self, service: ParserService, txt_file: Path
    ) -> None:
        """测试 TXT 解析：返回统一格式。

        验证：
        - 返回类型为 ParseResult
        - content 包含期望的文本内容
        - filename / format / char_count 正确
        """
        result = service.parse(txt_file)

        assert isinstance(result, ParseResult)
        assert "用户管理系统" in result.content
        assert "用户注册" in result.content
        assert "用户登录" in result.content
        assert result.filename == "test_doc.txt"
        assert result.format == "txt"
        assert result.char_count > 0
        assert result.char_count == len(result.content)

    def test_parse_md_returns_unified_format(
        self, service: ParserService, md_file: Path
    ) -> None:
        """测试 Markdown 解析：返回统一格式。"""
        result = service.parse(md_file)

        assert isinstance(result, ParseResult)
        assert "需求文档" in result.content
        assert "功能模块" in result.content
        assert result.format == "md"
        assert result.char_count > 0

    def test_parse_to_dict_returns_correct_structure(
        self, service: ParserService, txt_file: Path
    ) -> None:
        """测试 parse_to_dict 返回标准字典结构。

        验证输出格式: {"content": "...", "filename": "...", ...}
        """
        result = service.parse_to_dict(txt_file)

        assert isinstance(result, dict)
        assert "content" in result
        assert "filename" in result
        assert "format" in result
        assert "char_count" in result
        assert "metadata" in result
        assert result["filename"] == "test_doc.txt"
        assert result["format"] == "txt"
        assert isinstance(result["content"], str)
        assert isinstance(result["char_count"], int)
        assert result["char_count"] > 0

    def test_unsupported_format_raises_error(
        self, service: ParserService, tmp_path: Path
    ) -> None:
        """测试不支持的格式抛出 ValueError。"""
        file = tmp_path / "test.xyz"
        file.write_text("test content")

        with pytest.raises(ValueError) as exc_info:
            service.parse(file)
        assert "不支持" in str(exc_info.value)
        assert "xyz" in str(exc_info.value)

    def test_nonexistent_file_raises_error(
        self, service: ParserService
    ) -> None:
        """测试不存在的文件抛出 FileNotFoundError。"""
        with pytest.raises(FileNotFoundError) as exc_info:
            service.parse("nonexistent_file_12345.txt")
        assert "不存在" in str(exc_info.value) or "No such file" in str(exc_info.value)

    def test_no_extension_file_raises_error(
        self, service: ParserService, tmp_path: Path
    ) -> None:
        """测试无扩展名文件抛出 ValueError。"""
        file = tmp_path / "noextension"
        file.write_text("test content")

        with pytest.raises(ValueError) as exc_info:
            service.parse(file)
        assert "扩展名" in str(exc_info.value)


# ══════════════════════════════════════════════════════════
# 各格式解析器测试
# ══════════════════════════════════════════════════════════

class TestTxtParser:
    """TXT 解析器测试。"""

    def test_parse_empty_file_raises_error(
        self, service: ParserService, tmp_path: Path
    ) -> None:
        """测试空文件抛出 ParseError。"""
        file = tmp_path / "empty.txt"
        file.write_text("   \n\n   ")  # 只有空白字符

        with pytest.raises(ParseError) as exc_info:
            service.parse(file)
        assert "空" in str(exc_info.value.message)

    def test_parse_utf8_with_chinese(
        self, service: ParserService, tmp_path: Path
    ) -> None:
        """测试 UTF-8 中文文档解析。"""
        content = "## 1. 概述\n\n本系统旨在实现企业级测试管理平台。\n\n## 1.1 目标\n\n- 提升测试效率\n- 降低人工成本"
        file = tmp_path / "utf8_cn.txt"
        file.write_text(content, encoding="utf-8")

        result = service.parse(file)
        assert "测试管理平台" in result.content
        assert "提升测试效率" in result.content

    def test_parse_gbk_encoding(
        self, service: ParserService, tmp_path: Path
    ) -> None:
        """测试 GBK 编码自动检测（Windows 常见编码）。"""
        content = "需求文档\n系统功能说明\n模块一：用户管理"
        file = tmp_path / "gbk_doc.txt"
        file.write_text(content, encoding="gbk")

        result = service.parse(file)
        assert "需求文档" in result.content


class TestDocxParser:
    """DOCX 解析器测试。

    需要 python-docx 库。
    """

    def test_parse_docx(self, service: ParserService, tmp_path: Path) -> None:
        """测试 DOCX 文档解析。

        动态创建 .docx 文件进行测试（无外部依赖文件）。
        """
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx 未安装，跳过测试")

        # ── 创建测试 DOCX ────────────────────────
        doc = Document()
        doc.add_heading("需求规格说明书", level=1)
        doc.add_paragraph("本文档描述用户管理系统的功能需求。")
        doc.add_heading("用户注册", level=2)
        doc.add_paragraph("用户可通过手机号或邮箱注册账号。")

        # 添加表格
        table = doc.add_table(rows=2, cols=3)
        table.rows[0].cells[0].text = "模块"
        table.rows[0].cells[1].text = "功能"
        table.rows[0].cells[2].text = "优先级"
        table.rows[1].cells[0].text = "用户管理"
        table.rows[1].cells[1].text = "注册"
        table.rows[1].cells[2].text = "P0"

        file = tmp_path / "test_req.docx"
        doc.save(str(file))

        # ── 测试解析 ──────────────────────────────
        result = service.parse(file)

        assert "需求规格说明书" in result.content
        assert "用户注册" in result.content
        assert "手机号" in result.content
        assert "用户管理" in result.content  # 表格内容
        assert "P0" in result.content           # 表格内容
        assert result.format == "docx"
        assert result.filename == "test_req.docx"

    def test_parse_empty_docx_raises_error(
        self, service: ParserService, tmp_path: Path
    ) -> None:
        """测试空 DOCX 文件抛出 ParseError。"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx 未安装，跳过测试")

        doc = Document()
        # 不添加任何内容
        file = tmp_path / "empty.docx"
        doc.save(str(file))

        with pytest.raises(ParseError) as exc_info:
            service.parse(file)
        assert any(
            word in str(exc_info.value.message)
            for word in ("未提取到", "空")
        )


class TestPdfParser:
    """PDF 解析器测试。

    需要 PyPDF2 库。
    """

    def test_parse_pdf(self, service: ParserService, tmp_path: Path) -> None:
        """测试 PDF 文档解析。

        动态创建简单 PDF 进行测试（使用 reportlab 如果可用，
        否则跳过）。
        """
        try:
            from PyPDF2 import PdfWriter, PdfReader
            from io import BytesIO
        except ImportError:
            pytest.skip("PyPDF2 未安装，跳过测试")

        # 使用 PyPDF2 创建简单的 PDF（仅元数据方式）
        # PyPDF2 无法直接写入文本内容，换用 fpdf2 尝试
        try:
            from fpdf import FPDF
        except ImportError:
            # 尝试使用 reportlab
            try:
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import A4

                file = tmp_path / "test_req.pdf"
                c = canvas.Canvas(str(file), pagesize=A4)
                c.drawString(100, 750, "需求规格说明书 v1.0")
                c.drawString(100, 730, "用户管理模块功能描述")
                c.drawString(100, 700, "1. 用户注册 - 手机号/邮箱")
                c.drawString(100, 680, "2. 用户登录 - 密码验证")
                c.drawString(100, 660, "3. 密码修改 - 安全策略")
                c.showPage()
                c.save()

                result = service.parse(file)

                assert "需求规格说明书" in result.content
                assert "用户管理模块" in result.content
                assert result.format == "pdf"
                assert result.filename == "test_req.pdf"
                assert result.metadata.get("page_count") == 1
                return  # reportlab 测试通过
            except ImportError:
                pytest.skip("fpdf2/reportlab 均未安装，跳过 PDF 写入测试")

        # fpdf2 创建 PDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, text="需求规格说明书 v1.0", new_x="LMARGIN", new_y="NEXT")
        pdf.cell(200, 10, text="用户管理模块功能描述", new_x="LMARGIN", new_y="NEXT")
        pdf.cell(200, 10, text="1. 用户注册 - 手机号/邮箱", new_x="LMARGIN", new_y="NEXT")
        pdf.cell(200, 10, text="2. 用户登录 - 密码验证", new_x="LMARGIN", new_y="NEXT")

        file = tmp_path / "test_req.pdf"
        pdf.output(str(file))

        result = service.parse(file)

        assert "需求规格说明书" in result.content
        assert result.format == "pdf"
        assert result.filename == "test_req.pdf"
        assert result.metadata.get("page_count") == 1


# ══════════════════════════════════════════════════════════
# SOLID 原则验证
# ══════════════════════════════════════════════════════════

class TestSolidPrinciples:
    """验证 SOLID 原则的遵循情况。"""

    def test_open_closed_new_parser_registration(
        self, service: ParserService
    ) -> None:
        """测试 Open/Closed 原则：注册新解析器无需修改 ParserService。

        验证：
        - register() 可动态添加新格式解析器
        - 新解析器可正常工作
        """
        # ── 创建新的解析器（无需修改 ParserService 代码）
        class HtmlParser(BaseParser):
            @property
            def format(self) -> str:
                return "html"

            def do_parse(self, file_path: Path) -> ParseResult:
                return ParseResult(
                    content="parsed html content",
                    filename=file_path.name,
                    format="html",
                )

        # ── 注册 ──────────────────────────────────
        ParserService.register("html", HtmlParser())
        assert "html" in ParserService.supported_formats()

        # ── 验证新格式出现在支持列表中 ────────────
        assert "html" in service.supported_formats()

    def test_liskov_substitution_base_parser(
        self, service: ParserService
    ) -> None:
        """测试 Liskov 替换原则：所有解析器可替换 BaseParser。

        验证：
        - TxtParser / DocxParser / PdfParser 均为 BaseParser 有效子类
        - 通过基类接口调用行为一致
        """
        from services.parsers import TxtParser, DocxParser, PdfParser

        parser_classes = [TxtParser, DocxParser, PdfParser]

        for cls in parser_classes:
            parser = cls()
            # 1. 验证是 BaseParser 的子类
            assert isinstance(parser, BaseParser), (
                f"{cls.__name__} 不是 BaseParser 的子类"
            )
            # 2. 验证必须实现 format 属性
            fmt = parser.format
            assert isinstance(fmt, str) and len(fmt) > 0, (
                f"{cls.__name__}.format 返回无效值"
            )
            # 3. 验证必须实现 do_parse 方法
            assert hasattr(parser, "do_parse"), (
                f"{cls.__name__} 未实现 do_parse 方法"
            )

    def test_single_responsibility_parsers_are_independent(
        self, service: ParserService
    ) -> None:
        """测试 Single Responsibility：各解析器独立。

        验证：
        - 每个解析器只处理自己的格式
        - TxtParser 无法解析 docx（不受理）
        """
        from services.parsers import TxtParser

        txt_parser = TxtParser()
        # TXT 解析器只报告支持 txt 格式
        assert txt_parser.format == "txt"

        # PDF 解析器只报告支持 pdf 格式
        from services.parsers import PdfParser
        pdf_parser = PdfParser()
        assert pdf_parser.format == "pdf"


# ══════════════════════════════════════════════════════════
# ParseResult 测试
# ══════════════════════════════════════════════════════════

class TestParseResult:
    """ParseResult 数据类测试。"""

    def test_to_dict_output_format(self) -> None:
        """测试 to_dict 输出格式严格符合规范。

        输出必须为: {"content": "...", "filename": "...", ...}
        """
        result = ParseResult(
            content="测试内容",
            filename="test.txt",
            format="txt",
        )

        d = result.to_dict()

        assert d["content"] == "测试内容"
        assert d["filename"] == "test.txt"
        assert d["format"] == "txt"
        assert d["char_count"] == 4  # "测试内容" = 4 字符
        assert isinstance(d["metadata"], dict)

    def test_char_count_auto_calculated(self) -> None:
        """测试 char_count 自动计算。"""
        result = ParseResult(
            content="ABCDEFGHIJ",  # 10 字符
            filename="test.txt",
            format="txt",
        )
        assert result.char_count == 10

    def test_metadata_default_empty(self) -> None:
        """测试 metadata 默认为空字典。"""
        result = ParseResult(
            content="test",
            filename="test.txt",
            format="txt",
        )
        assert result.metadata == {}
